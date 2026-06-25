from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_doc():
    doc = Document()
    
    # Title
    title = doc.add_heading('Lampiran (opsional)', level=1)
    
    doc.add_heading('Potongan Kode Penting', level=2)
    doc.add_paragraph('Berikut adalah potongan kode yang merepresentasikan inti pemrosesan data, yaitu perbedaan implementasi pencarian antara Vector dan Unordered Map, serta cara perhitungan estimasi penggunaan memori.')
    
    doc.add_paragraph('1. Implementasi Pencarian menggunakan std::vector (Linear Search O(n))').bold = True
    code1 = doc.add_paragraph('''vector<TransactionItem> searchByTransaction(const string& transactionId) const {
    vector<TransactionItem> result;
    // Harus melakukan iterasi linier pada seluruh elemen
    for (const auto& item : transactions) {
        if (item.transactionId == transactionId) result.push_back(item);
    }
    return result;
}''')
    code1.style = 'Quote'

    doc.add_paragraph('2. Implementasi Pencarian menggunakan std::unordered_map (Hash Lookup O(1))').bold = True
    code2 = doc.add_paragraph('''vector<TransactionItem> searchByTransaction(const string& transactionId) const {
    vector<TransactionItem> result;
    // Menggunakan indeks hash untuk langsung melompat ke posisi data
    auto it = indexByTransaction.find(transactionId);
    if (it == indexByTransaction.end()) return result;

    for (int index : it->second) result.push_back(transactions[index]);
    return result;
}''')
    code2.style = 'Quote'
    
    doc.add_paragraph('3. Perhitungan Estimasi Memori').bold = True
    code3 = doc.add_paragraph('''size_t estimateMemoryBytes() const {
    // Menghitung ukuran container + alokasi bucket di Unordered Map
    size_t memory = transactions.size() * sizeof(TransactionItem) + products.size() * sizeof(Product);
    memory += indexByTransaction.size() * sizeof(pair<const string, vector<int>>);
    memory += indexByCustomer.size() * sizeof(pair<const string, vector<int>>);
    memory += indexByProduct.size() * sizeof(pair<const string, vector<int>>);
    return memory;
}''')
    code3.style = 'Quote'

    doc.add_heading('Data Tambahan Eksperimen', level=2)
    doc.add_paragraph('Bagian ini memuat sampel struktur dataset yang digunakan selama proses pengujian eksperimental serta hasil tangkapan layar (screenshot) dari implementasi operasi CRUD secara komparatif.')
    
    doc.add_paragraph('A. Sampel Struktur Dataset CSV').bold = True
    doc.add_paragraph('products_300.csv: Menyimpan ID produk, nama, dan kategorinya.')
    doc.add_paragraph('P001,Produk_1,Elektronik\\nP002,Produk_2,Aksesoris\\nP003,Produk_3,Komputer', style='Quote')
    
    doc.add_paragraph('transactions-2000.csv: Menyimpan relasi ID transaksi, pelanggan, tanggal, relasi ID produk, dan kuantitas.')
    doc.add_paragraph('T00001,C001,2026-05-01,P001,Produk_1,Elektronik,1\\nT00001,C001,2026-05-01,P002,Produk_2,Aksesoris,2\\nT00001,C001,2026-05-01,P003,Produk_3,Komputer,3', style='Quote')

    doc.add_paragraph('B. Dokumentasi Hasil Operasi CRUD pada Sistem').bold = True
    
    doc.add_paragraph('1. Operasi Insert Data:')
    if os.path.exists('screenshot_insert.jpg'):
        doc.add_picture('screenshot_insert.jpg', width=Inches(6.0))
        
    doc.add_paragraph('2. Operasi Search Data:')
    if os.path.exists('screenshot_search.jpg'):
        doc.add_picture('screenshot_search.jpg', width=Inches(6.0))
        
    doc.add_paragraph('3. Operasi Update Data:')
    if os.path.exists('screenshot_update.jpg'):
        doc.add_picture('screenshot_update.jpg', width=Inches(6.0))
        
    doc.add_paragraph('4. Operasi Delete Data:')
    if os.path.exists('screenshot_delete.jpg'):
        doc.add_picture('screenshot_delete.jpg', width=Inches(6.0))

    doc.add_paragraph('C. Output Pengujian Benchmark').bold = True
    doc.add_paragraph('Berikut adalah hasil dari eksekusi program benchmark yang membandingkan performa antara Vector dan Unordered Map pada dataset besar secara komprehensif.')
    
    doc.add_paragraph('Tabel Komparasi Sebelahan (Side-by-Side Comparison):')
    if os.path.exists('terminal_lampiran_c.jpg'):
        doc.add_picture('terminal_lampiran_c.jpg', width=Inches(7.0))
        
    doc.add_paragraph('Output Evaluasi Waktu, Memori, dan Frequently Bought Together:')
    if os.path.exists('terminal_lampiran_b.jpg'):
        doc.add_picture('terminal_lampiran_b.jpg', width=Inches(6.0))

    doc.save('Lampiran_Opsional_Final_Lengkap.docx')
    print('Docx Lampiran Opsional updated with benchmark.')

if __name__ == '__main__':
    create_doc()
