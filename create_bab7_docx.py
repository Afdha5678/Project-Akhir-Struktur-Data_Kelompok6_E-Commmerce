from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_doc():
    doc = Document()
    
    # Title
    title = doc.add_heading('BAB VII\nKESIMPULAN DAN REKOMENDASI', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading('7.1 Ringkasan Temuan Utama', level=2)
    doc.add_paragraph('Berdasarkan pengujian performa yang telah dilakukan melalui program benchmark (benchmark (2).cpp) pada dua struktur data utama, yaitu std::vector dan std::unordered_map (Hash Table) menggunakan dataset simulasi transaksi masif, ditemukan beberapa poin kunci:')
    
    doc.add_paragraph('1. Kecepatan Waktu Eksekusi (Time Complexity):', style='List Number')
    p = doc.add_paragraph('Struktur data ')
    p.add_run('Unordered Map').bold = True
    p.add_run(' menunjukkan performa waktu yang jauh lebih unggul pada seluruh operasi pencarian (Search by Transaction, Customer, Product) dan analitik rekomendasi (Top-N Products, Frequently Bought Together). Unordered Map mengeksekusi operasi tersebut dalam waktu hampir konstan O(1) rata-rata. Sebaliknya, struktur data ')
    p.add_run('Vector').bold = True
    p.add_run(' mengalami penurunan performa pencarian yang melambat secara linear O(n) seiring bertambahnya jumlah data transaksi.')
    
    doc.add_paragraph('Bukti Kode Pencarian O(1) pada Unordered Map:').bold = True
    code1 = doc.add_paragraph('''vector<TransactionItem> searchByTransaction(const string& transactionId) {
    vector<TransactionItem> result;
    // Pencarian langsung ke indeks hash table (O(1))
    auto it = indexByTransaction.find(transactionId);
    if (it == indexByTransaction.end()) return result;

    for (int idx : it->second) result.push_back(transactions[idx]);
    return result;
}''')
    code1.style = 'Quote'
    
    doc.add_paragraph('2. Penggunaan Memori (Space Complexity):', style='List Number')
    p2 = doc.add_paragraph('Walaupun Vector lebih lambat dalam eksekusi pencarian, struktur memori ini terbukti ')
    p2.add_run('lebih hemat memori').bold = True
    p2.add_run(' dibandingkan Unordered Map. Vector menyimpan objek secara berurutan (contiguous) tanpa perlu referensi ekstra, sementara Unordered Map harus menyimpan pointer/index untuk berbagai kategori (Transaction, Customer, Product) demi membentuk arsitektur hash table, yang memakan ukuran memori sekitar 60%-80% lebih besar dari Vector.')
    
    doc.add_paragraph('Ilustrasi Hasil Pencarian & Perbandingan:').bold = True
    if os.path.exists('screenshot_search.jpg'):
        doc.add_picture('screenshot_search.jpg', width=Inches(6.0))
        
    doc.add_heading('7.2 Struktur Data Paling Optimal dan Alasannya', level=2)
    p3 = doc.add_paragraph('Untuk pengembangan ')
    p3.add_run('Sistem Rekomendasi Produk Berbasis Riwayat Transaksi').bold = True
    p3.add_run(', struktur data yang paling optimal adalah ')
    p3.add_run('Unordered Map').bold = True
    p3.add_run('.')
    
    doc.add_paragraph('Alasan Utama:').bold = True
    doc.add_paragraph('Skalabilitas Kecepatan Pencarian: Dalam simulasi dunia nyata, platform e-commerce memproses jutaan histori pencarian tiap harinya. Menggunakan metode linear search pada Vector tidak mungkin mengakomodasi kebutuhan response-time cepat. Oleh karena itu, kemampuan std::unordered_map yang memproses query dalam hitungan skala mikrosekon O(1) sangat mutlak diperlukan demi kenyamanan pengguna.', style='List Number')
    doc.add_paragraph('Keunggulan Sistem Rekomendasi Berulang: Saat program memproses algoritma Frequently Bought Together, sistem harus menghubungkan ratusan bahkan ribuan transaksi beririsan. Penggunaan Vector akan mengakibatkan proses iterasi eksponensial bersarang (nested loops), sementara unordered_map mampu melakukan lompatan indeks secara efisien.', style='List Number')
    doc.add_paragraph('Trade-Off Ruang dan Waktu: Biaya penambahan beban Memory (RAM) pada Unordered Map dinilai merupakan kompensasi (trade-off) yang wajar dan sepadan demi meminimalisasi latency pemrosesan data (Waktu).', style='List Number')

    doc.add_paragraph('Cuplikan Program Analisis Data (Frequently Bought Together):').bold = True
    code2 = doc.add_paragraph('''void displayFrequentlyBoughtTogether() {
    // Penggunaan unordered_set dan unordered_map untuk agregasi
    unordered_set<string> relatedTransactionIds;
    for (int idx : productIndexIt->second) {
        relatedTransactionIds.insert(transactions[idx].transactionId);
    }

    unordered_map<string, int> togetherCounts;
    for (const string& transactionId : relatedTransactionIds) {
        auto transactionIt = indexByTransaction.find(transactionId);
        // ... (Agregasi frekuensi produk komplementer)
}''')
    code2.style = 'Quote'

    doc.add_paragraph('Ilustrasi Hasil Analisis Benchmark Waktu:').bold = True
    if os.path.exists('table_waktu.jpg'):
        doc.add_picture('table_waktu.jpg', width=Inches(5.0))
        
    doc.add_heading('7.3 Saran Pengembangan Lanjutan', level=2)
    doc.add_paragraph('Berdasarkan implementasi sistem rekomendasi yang telah dikerjakan, diusulkan beberapa rekomendasi untuk pengembangan tahap selanjutnya:')
    
    doc.add_paragraph('Penerapan Struktur Data Graph (Graf): Sebagai alternatif atau pendamping Hash Map, sistem rekomendasi produk dapat memanfaatkan struktur Graf (seperti Adjacency List) di masa depan. Algoritma Graf jauh lebih fleksibel dalam memetakan serta menavigasi korelasi many-to-many antar produk (nodes) maupun relasi kedekatannya (edges/weights).', style='List Number')
    doc.add_paragraph('Peralihan ke Database Relasional Modern (RDBMS): Mengingat saat ini eksekusi I/O pada products_300.csv dan transactions-2000.csv dibebankan sepenuhnya ke beban kerja memori, sangat disarankan agar kedepannya sistem mengintegrasikan database seperti PostgreSQL. RDBMS telah memiliki kemampuan indexing (B-Tree/Hash) bawaan yang siap menangani relasi antar-tabel dengan lebih stabil tanpa berisiko data korup (data corruption).', style='List Number')
    doc.add_paragraph('Implementasi Pagination (Paging): Saat jumlah output sangat besar, disarankan penambahan antarmuka pagination pada terminal untuk memudahkan pengguna dalam membaca baris-baris rekomendasi tanpa perlu melakukan gulir (scrolling) panjang ke atas.', style='List Number')
    
    doc.save('Bab_VII_Kesimpulan_Rekomendasi.docx')

if __name__ == "__main__":
    create_doc()
