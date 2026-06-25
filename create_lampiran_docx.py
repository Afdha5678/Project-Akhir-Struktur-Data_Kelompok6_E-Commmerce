from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_doc():
    doc = Document()
    
    # Title
    title = doc.add_heading('Lampiran Benchmark', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('Berikut adalah hasil dari eksekusi benchmark membandingkan Vector vs Unordered Map pada dataset 2000 baris:')
    
    doc.add_heading('1. Tabel Komparasi Sebelahan (Side-by-Side Comparison)', level=2)
    if os.path.exists('terminal_lampiran_c.jpg'):
        doc.add_picture('terminal_lampiran_c.jpg', width=Inches(7.0))
    else:
        doc.add_paragraph('[Gambar terminal_lampiran_c.jpg tidak ditemukan]')

    doc.add_heading('2. Output Evaluasi Waktu, Memori, dan Frequently Bought Together', level=2)
    if os.path.exists('terminal_lampiran_b.jpg'):
        doc.add_picture('terminal_lampiran_b.jpg', width=Inches(6.0))
    else:
        doc.add_paragraph('[Gambar terminal_lampiran_b.jpg tidak ditemukan]')
        
    doc.save('Lampiran_Benchmark_Full.docx')
    print('Docx generated.')

if __name__ == '__main__':
    create_doc()
